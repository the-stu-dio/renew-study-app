from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt, ensure_csrf_cookie
import json
import traceback
import uuid
import os
import io
from django.template.loader import render_to_string

def landing_page_view(request):
    """Landing page - first page visitors see"""
    return render(request, 'accounts/landing.html')

def intro_video_1_view(request):
    """First intro video page"""
    return render(request, 'accounts/intro_video_1.html')

def intro_video_2_view(request):
    """Second intro video page"""
    return render(request, 'accounts/intro_video_2.html')

def nickname_view(request):
    """Set user nickname - stored in localStorage"""
    return render(request, 'accounts/nickname.html')

def welcome_view(request):
    """Welcome page after setting nickname"""
    return render(request, 'accounts/welcome.html')

def hub_view(request):
    """Main hub page"""
    return render(request, 'accounts/hub.html')

def review_renew_view(request):
    """Review RENEW landing page - shows small boxes for quick review"""
    return render(request, 'accounts/review_renew.html')

def profile_view(request):
    """User profile page - data loaded from localStorage"""
    return render(request, 'accounts/profile.html')

def adventure_map_view(request):
    """Adventure map showing progress"""
    return render(request, 'accounts/adventure_map.html')

# Positive Events Views
def pos_p1_view(request):
    """Positive Events Part 1"""
    return render(request, 'accounts/pos_p1.html')

def pos_p2_view(request):
    """Positive Events Part 2"""
    return render(request, 'accounts/pos_p2.html')

def pos_p3_view(request):
    """Positive Events Part 3"""
    return render(request, 'accounts/pos_p3.html')

def pos_p4_view(request):
    """Positive Events Part 4"""
    return render(request, 'accounts/pos_p4.html')
    
def pos_p5_view(request):
    """Positive Events Part 5 - Interactive Cards"""
    return render(request, 'accounts/pos_p5.html')
    
# Negative Events Views
def neg_p1_view(request):
    """Negative Events Part 1"""
    return render(request, 'accounts/neg_p1.html')

def neg_p2_view(request):
    """Negative Events Part 2"""
    return render(request, 'accounts/neg_p2.html')

def neg_p3_view(request):
    """Negative Events Part 3"""
    return render(request, 'accounts/neg_p3.html')

def neg_p4_view(request):
    """Negative Events Part 4 - Interactive Cards"""
    return render(request, 'accounts/neg_p4.html')

def neg_p5_view(request):
    """Negative Events Part 5"""
    return render(request, 'accounts/neg_p5.html')

def neg_p6_view(request):
    """Negative Events Part 6"""
    return render(request, 'accounts/neg_p6.html')

def neg_p7_view(request):
    """Negative Events Part 7"""
    return render(request, 'accounts/neg_p7.html')

def neg_p8_view(request):
    """Negative Events Part 8"""
    return render(request, 'accounts/neg_p8.html')

def neg_p9_view(request):
    """Negative Events Part 9"""
    return render(request, 'accounts/neg_p9.html')

def neg_p10_view(request):
    """Negative Events Part 10 - Quiz"""
    return render(request, 'accounts/neg_p10.html')

# Thoughts & Emotions Views
def emo_p1_view(request):
    """Thoughts and Feelings Part 1"""
    return render(request, 'accounts/emo_p1.html')

def emo_p2_view(request):
    """Thoughts and Feelings Part 2"""
    return render(request, 'accounts/emo_p2.html')

def emo_p3_view(request):
    """Thoughts and Feelings Part 3"""
    return render(request, 'accounts/emo_p3.html')

def emo_p4_view(request):
    """Thoughts and Feelings Part 4"""
    return render(request, 'accounts/emo_p4.html')

def emo_p5_view(request):
    """Thoughts and Feelings Part 5"""
    return render(request, 'accounts/emo_p5.html')

# Skills Views
def skill_p1_view(request):
    """Skills Part 1"""
    return render(request, 'accounts/skill_p1.html')

def skill_p2_view(request):
    """Skills Part 2"""
    return render(request, 'accounts/skill_p2.html')
    
def skill_p3_view(request):
    """Skills Part 3"""
    return render(request, 'accounts/skill_p3.html')

@ensure_csrf_cookie
def skill_p4_view(request):
    """Skills Part 4 - Three Skill Categories Overview"""
    return render(request, 'accounts/skill_p4.html')

def journal_p1_view(request):
    """Journaling Part 1"""
    return render(request, 'accounts/journal_p1.html')

def journal_p2_view(request):
    """Journaling Part 2"""
    return render(request, 'accounts/journal_p2.html')

def journal_option1_view(request):
    """Journaling Option 1 - Things I Want to Work On"""
    return render(request, 'accounts/journal_option1.html')

def journal_option2_view(request):
    """Journaling Option 2 - Experiences That Shaped Me"""
    return render(request, 'accounts/journal_option2.html')

def posplan_p1_view(request):
    """Positive Planned Activities Part 1"""
    return render(request, 'accounts/posplan_p1.html')

def posplan_p2_view(request):
    """Positive Planned Activities Part 2 - Who"""
    return render(request, 'accounts/posplan_p2.html')

def posplan_p3_view(request):
    """Positive Planned Activities Part 3 - Where"""
    return render(request, 'accounts/posplan_p3.html')

def posplan_p4_view(request):
    """Positive Planned Activities Part 4 - What Activity"""
    return render(request, 'accounts/posplan_p4.html')

def posplan_p5_view(request):
    """Positive Planned Activities Part 5 - When/Timing"""
    return render(request, 'accounts/posplan_p5.html')

def posplan_summary_p1_view(request):
    """Positive Planned Activities - Summary Part 1"""
    return render(request, 'accounts/posplan_summary_p1.html')
    
def posplan_summary_p2_view(request):
    """Positive Planned Activities - Summary Part 2"""
    return render(request, 'accounts/posplan_summary_p2.html')

def posplan_summary_p3_view(request):
    """Positive Planned Activities - Summary Part 3"""
    return render(request, 'accounts/posplan_summary_p3.html')

def posplan_summary_p4_view(request):
    """Positive Planned Activities - Summary Part 4"""
    return render(request, 'accounts/posplan_summary_p4.html')

def pmr_p1_view(request):
    """Progressive Muscle Relaxation Part 1"""
    return render(request, 'accounts/pmr_p1.html')

def discoveries_view(request):
    """My Discoveries page - data loaded from localStorage"""
    return render(request, 'accounts/discoveries.html')

def tips(request):
    return render(request, 'accounts/tips.html')


def _purge_other_sessions(current_session_key):
    """Delete all response rows that belong to a DIFFERENT session.
    This keeps the DB clean: only one participant's data exists at a time.
    """
    from accounts.models import (
        PositiveEventResponse, NegativeEventResponse,
        PMRSurveyResponse, AttentionCheckResponse, IronSpongeResponse,
        PosplanActivityResponse,
    )
    for Model in [
        PositiveEventResponse, NegativeEventResponse,
        PMRSurveyResponse, AttentionCheckResponse, IronSpongeResponse,
        PosplanActivityResponse,
    ]:
        Model.objects.exclude(session_key=current_session_key).delete()


@csrf_exempt
def save_positive_response(request):
    """Save a positive event response to the database"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            category = data.get('category', '')
            response_text = data.get('response_text', '')
            nickname = data.get('nickname', '')

            if not category or not response_text:
                return JsonResponse({'error': 'Category and response_text are required'}, status=400)

            valid_categories = ['recreation', 'achievement', 'relationship', 'other']
            if category not in valid_categories:
                return JsonResponse({'error': f'Invalid category. Must be one of: {valid_categories}'}, status=400)

            # Ensure a session exists for anonymous users
            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            # Purge data from any previous session so only this participant exists
            _purge_other_sessions(session_key)

            from accounts.models import PositiveEventResponse
            obj, created = PositiveEventResponse.objects.update_or_create(
                session_key=session_key,
                category=category,
                defaults={
                    'response_text': response_text,
                    'nickname': nickname,
                    'user': request.user if request.user.is_authenticated else None,
                }
            )

            return JsonResponse({
                'success': True,
                'created': created,
                'category': category,
            })
        except Exception as e:
            # Print full traceback to server console for debugging
            print('submit_feedback_survey: exception:\n', traceback.format_exc())
            return JsonResponse({'error': str(e), 'trace': traceback.format_exc()}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)

@csrf_exempt
def save_negative_response(request):
    """Save a negative event response to the database"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            category = data.get('category', '')
            response_text = data.get('response_text', '')
            nickname = data.get('nickname', '')

            if not category or not response_text:
                return JsonResponse({'error': 'Category and response_text are required'}, status=400)

            valid_categories = ['adverse', 'neighbor', 'discriminate', 'stress']
            if category not in valid_categories:
                return JsonResponse({'error': f'Invalid category. Must be one of: {valid_categories}'}, status=400)

            # Ensure a session exists for anonymous users
            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            _purge_other_sessions(session_key)

            from accounts.models import NegativeEventResponse
            obj, created = NegativeEventResponse.objects.update_or_create(
                session_key=session_key,
                category=category,
                defaults={
                    'response_text': response_text,
                    'nickname': nickname,
                    'user': request.user if request.user.is_authenticated else None,
                }
            )

            return JsonResponse({
                'success': True,
                'created': created,
                'category': category,
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)

# Optional: Keep these endpoints for backwards compatibility, but they won't be used
@csrf_exempt
def mark_video_complete(request):
    """AJAX endpoint - now handled by localStorage, but keeping for compatibility"""
    if request.method == 'POST':
        return JsonResponse({'success': True, 'message': 'Handled by localStorage'})
    return JsonResponse({'error': 'Invalid method'}, status=405)

@csrf_exempt
def submit_feedback_survey(request):
    """Accept feedback survey responses (emo_p4) and save to DB (session-keyed).

    Expects JSON POST: { "checked": ["learned_about_world_like_talking", ...], "nickname": "..." }
    Saves a FeedbackSurveyResponse keyed by session_key so the summary generator can include it.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body or b'{}')
            checked = data.get('checked', []) or []
            nickname = data.get('nickname', '')

            # Map incoming option values to model boolean fields
            mapping = {
                'learned_about_world_like_talking': 'statement1_checked',
                'learned_about_world_want_more': 'statement2_checked',
                'found_some_boring': 'statement3_checked',
                'didnt_find_helpful': 'statement4_checked',
            }

            # Human-readable statement text mapping
            text_mapping = {
                'learned_about_world_like_talking': 'I liked learning about how the world around us affects people, and I want to talk to someone more about this.',
                'learned_about_world_want_more': 'I liked learning about how the world around us affects people, and I want to see or read more about it.',
                'found_some_boring': 'To be honest, I found some of this stuff a little boring, but I want to talk more about thoughts, feelings, and emotions.',
                'didnt_find_helpful': 'To be honest, I didn\'t find what we talked about so far very interesting or helpful.',
            }

            # Ensure a session exists
            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            _purge_other_sessions(session_key)

            from accounts.models import FeedbackSurveyResponse

            defaults = {
                'nickname': nickname,
                'statement1_checked': False,
                'statement2_checked': False,
                'statement3_checked': False,
                'statement4_checked': False,
                'responses': [],
            }

            for opt in checked:
                field = mapping.get(opt)
                if field:
                    defaults[field] = True
            # Build list of actual statement texts for storage
            selected_texts = [text_mapping.get(opt) for opt in checked if text_mapping.get(opt)]
            defaults['responses'] = selected_texts

            # Debug logging: print incoming payload and session
            print('submit_feedback_survey: received checked=', checked, 'nickname=', nickname, 'session_key=', session_key)

            obj, created = FeedbackSurveyResponse.objects.update_or_create(
                session_key=session_key,
                defaults=defaults
            )

            print(f'submit_feedback_survey: saved FeedbackSurveyResponse id={obj.id} created={created}')

            return JsonResponse({'success': True, 'created': created, 'session_key': session_key, 'id': obj.id})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)

def submit_journaling_survey(request):
    """Accept journaling survey responses and persist to FeedbackSurveyResponse.responses

    Expects JSON: { "responses": ["share_family", ...], "nickname": "..." }
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body or b'{}')
            selected = data.get('responses', []) or []
            nickname = data.get('nickname', '')

            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            # Map the data-value codes to human-readable labels (matches skill_p4.html)
            mapping = {
                'share_family': 'I like writing about these things, and I want to share it with a family member or friend.',
                'share_counselor': 'I like writing about these things, and I want to share it with a counselor.',
                'like_private': "I like writing about these things, but I don't want to share with anyone.",
                'prefer_talking': "I don't like writing about these things. I would rather talk to someone.",
                'dislike_both': "I don't like writing or talking about these things. It makes me feel worse.",
            }

            labels = [mapping.get(code, code) for code in selected]

            from accounts.models import FeedbackSurveyResponse

            obj, created = FeedbackSurveyResponse.objects.get_or_create(
                session_key=session_key,
                defaults={'nickname': nickname, 'responses': []}
            )

            existing = obj.responses or []
            for lbl in labels:
                if lbl and lbl not in existing:
                    existing.append(lbl)

            obj.responses = existing
            if nickname:
                obj.nickname = nickname
            obj.save()

            return JsonResponse({'success': True, 'saved': labels, 'created': created})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)

@csrf_exempt
def submit_pmr_survey(request):
    """Save PMR survey response to the database"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            selected_feelings = data.get('selected_feelings', [])
            nickname = data.get('nickname', '')

            if not selected_feelings:
                return JsonResponse({'error': 'selected_feelings is required'}, status=400)

            valid_feelings = ['relaxed', 'stressed', 'same']
            for f in selected_feelings:
                if f not in valid_feelings:
                    return JsonResponse({'error': f'Invalid feeling: {f}. Must be one of: {valid_feelings}'}, status=400)

            # Ensure a session exists for anonymous users
            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            _purge_other_sessions(session_key)

            from accounts.models import PMRSurveyResponse
            obj, created = PMRSurveyResponse.objects.update_or_create(
                session_key=session_key,
                defaults={
                    'selected_feelings': selected_feelings,
                    'nickname': nickname,
                    'user': request.user if request.user.is_authenticated else None,
                }
            )

            return JsonResponse({
                'success': True,
                'created': created,
                'selected_feelings': selected_feelings,
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)


@csrf_exempt
def submit_pmr_full(request):
    """Save the full PMR survey selections (main modal) to FeedbackSurveyResponse.responses
    Expects JSON: { "selected_options": [...], "nickname": "..." }
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body or b'{}')
            selected = data.get('selected_options', []) or []
            nickname = data.get('nickname', '')

            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            # Map option codes to human-readable labels (should match pmr_p1.js texts)
            mapping = {
                'category_liked': 'I liked it',
                'sub_liked_own': 'I want to try it on my own',
                'sub_liked_counselor': 'I would like to try doing it more with a counselor',
                'category_didnt_like': "I didn't like it",
                'sub_didnt_like_no_again': 'I would prefer not to try something like that again',
                'sub_didnt_like_try_counselor': 'I would be open to trying it with a counselor again to see if there was something I was doing wrong',
                'sub_didnt_like_other_relaxation': 'I would be interested in talking to a counselor about other things I can practice to feel more relaxed',
                'category_dont_know': "I don't know how I feel about it",
            }

            labels = [mapping.get(code, code) for code in selected]

            from accounts.models import FeedbackSurveyResponse

            # Merge into existing FeedbackSurveyResponse.responses (avoid duplicates)
            obj, created = FeedbackSurveyResponse.objects.get_or_create(
                session_key=session_key,
                defaults={'nickname': nickname, 'responses': []}
            )

            existing = obj.responses or []
            # Add new labels preserving order, avoid dupes
            for lbl in labels:
                if lbl and lbl not in existing:
                    existing.append(lbl)

            obj.responses = existing
            if nickname:
                obj.nickname = nickname
            obj.save()

            return JsonResponse({'success': True, 'saved': labels, 'created': created})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)


@csrf_exempt
def submit_posplan_survey(request):
    """Save Positive Activity Plan feeling selections into FeedbackSurveyResponse.responses

    Expects JSON: { "selected": ["will_happen","excited",...], "nickname": "..." }
    Maps codes to human-readable labels and appends to FeedbackSurveyResponse.responses.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body or b'{}')
            selected = data.get('selected', []) or []
            nickname = data.get('nickname', '')

            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            # Map codes to labels (must match UI texts)
            mapping = {
                'will_happen': 'I think this positive planned activity will happen',
                'excited': "I'm excited",
                'counselor_help': "But I'd love to talk to a counselor about it to make sure it happens!",
                'wont_happen': "I don't think this will happen",
                'not_sure': "I'm not sure this will happen",
                'counselor_help_make': "But I'd like to talk to a counselor to give me help to make it happen",
                'someone_else': 'I would like to share it with someone else at school',
                'no_sharing': "I don't want to share this with anyone",
                'open_counselor': "But I'd be open to talking to a counselor about it",
                'no_counselor': "I don't want to talk to a counselor about it",
            }

            labels = [mapping.get(code, code) for code in selected]

            from accounts.models import FeedbackSurveyResponse

            obj, created = FeedbackSurveyResponse.objects.get_or_create(
                session_key=session_key,
                defaults={'nickname': nickname, 'responses': []}
            )

            existing = obj.responses or []
            for lbl in labels:
                if lbl and lbl not in existing:
                    existing.append(lbl)

            obj.responses = existing
            if nickname:
                obj.nickname = nickname
            obj.save()

            return JsonResponse({'success': True, 'saved': labels, 'created': created})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)

@csrf_exempt
def save_attention_check(request):
    """Track incorrect attempts on attention check questions"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            question_id = data.get('question_id', '')
            is_correct = data.get('is_correct', False)
            nickname = data.get('nickname', '')

            if not question_id:
                return JsonResponse({'error': 'question_id is required'}, status=400)

            valid_questions = ['neg_p6_q1', 'neg_p7_q1', 'emo_p2_q1', 'emo_p2_q2', 'emo_p2_q3', 'emo_p2_q4']
            if question_id not in valid_questions:
                return JsonResponse({'error': f'Invalid question_id. Must be one of: {valid_questions}'}, status=400)

            # Ensure a session exists
            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            _purge_other_sessions(session_key)

            from accounts.models import AttentionCheckResponse
            obj, created = AttentionCheckResponse.objects.get_or_create(
                session_key=session_key,
                question_id=question_id,
                defaults={
                    'nickname': nickname,
                    'user': request.user if request.user.is_authenticated else None,
                }
            )

            if is_correct:
                obj.answered_correctly = True
            else:
                obj.incorrect_attempts += 1

            obj.nickname = nickname or obj.nickname
            obj.save()

            return JsonResponse({
                'success': True,
                'question_id': question_id,
                'incorrect_attempts': obj.incorrect_attempts,
                'answered_correctly': obj.answered_correctly,
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)

def get_posplan_summary(request):
    """AJAX endpoint - now handled by localStorage"""
    if request.method == 'GET':
        return JsonResponse({'success': True, 'message': 'Handled by localStorage'})
    return JsonResponse({'error': 'Invalid method'}, status=405)

@csrf_exempt
def save_journal_entry(request):
    """AJAX endpoint to save journal entries to the database."""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            session_key = request.session.session_key
            if not session_key:
                request.session.create()
                session_key = request.session.session_key

            from accounts.models import JournalEntry
            entry, created = JournalEntry.objects.get_or_create(
                session_key=session_key,
                defaults={'nickname': data.get('nickname', '')}
            )

            if 'things_to_change' in data:
                entry.things_to_change = data['things_to_change']
            if 'experiences_shaped_me' in data:
                entry.experiences_shaped_me = data['experiences_shaped_me']
            if 'nickname' in data:
                entry.nickname = data['nickname']

            entry.save()

            return JsonResponse({
                'success': True,
                'message': 'Journal entry saved',
                'session_key': session_key
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)

# New endpoint for Qualtrics export
# Replace your existing export_to_qualtrics function with this enhanced version

@csrf_exempt
def export_to_qualtrics(request):
    """Export data to Qualtrics - receives localStorage data from frontend"""
    if request.method == 'POST':
        try:
            # Receive the localStorage data from the frontend
            data = json.loads(request.body)
            
            # Generate a participant ID for tracking
            participant_id = str(uuid.uuid4())[:8]  # Short ID
            
            # Extract main renew_study_data
            renew_data = data.get('renew_study_data', {})
            
            # Extract ironSpongeResults
            iron_sponge_data = data.get('ironSpongeResults', {})
            
            # Process journal entries
            journal_entries = renew_data.get('journal_entries', {})
            journal_things_to_change = journal_entries.get('things_to_change', {}).get('content', '')
            
            # Process experiences journal (JSON format)
            experiences_content = journal_entries.get('experiences_shaped_me', {}).get('content', '')
            journal_answers = {}
            if experiences_content:
                try:
                    parsed_experiences = json.loads(experiences_content)
                    journal_answers = {
                        'answer1': parsed_experiences.get('answer1', ''),
                        'answer2': parsed_experiences.get('answer2', ''),
                        'answer3': parsed_experiences.get('answer3', ''),
                        'answer4': parsed_experiences.get('answer4', '')
                    }
                except json.JSONDecodeError:
                    journal_answers = {'raw_content': experiences_content}
            
            # Add things_to_change to journal_answers
            journal_answers['things_to_change'] = journal_things_to_change
            
            # Process posplan responses
            posplan_responses = renew_data.get('posplan_responses', {})
            
            # Process survey responses
            survey_responses = renew_data.get('survey_responses', {})
            feedback_survey = survey_responses.get('feedback_survey', {})
            journaling_survey = survey_responses.get('journaling_survey', {})
            pmr_survey = survey_responses.get('pmr_survey', {})
            pmr_sidebar_survey = survey_responses.get('pmr_sidebar_survey', {})
            posplan_survey = renew_data.get('posplan_survey', {})
            
            # Extract additional teaching survey data
            additional_teaching_survey = renew_data.get('additional_teaching_survey', {})
            individual_privacy_settings = renew_data.get('individual_privacy_settings', {})
            
            # Debug: Print all survey response keys to help identify missing data
            print("All survey response keys:", list(survey_responses.keys()))
            print("Additional teaching survey:", additional_teaching_survey)
            
            # Process Iron & Sponge results
            iron_sponge_summary = {}
            iron_sponge_questions = {}
            
            if iron_sponge_data:
                # Domain summaries
                for domain in ['school', 'family', 'friend']:
                    domain_data = iron_sponge_data.get(domain, {})
                    if domain_data:
                        iron_sponge_summary[f'{domain}_type'] = domain_data.get('type', '')
                        counts = domain_data.get('counts', {})
                        percentages = domain_data.get('percentages', {})
                        iron_sponge_summary[f'{domain}_iron_count'] = counts.get('iron', 0)
                        iron_sponge_summary[f'{domain}_sponge_count'] = counts.get('sponge', 0)
                        iron_sponge_summary[f'{domain}_neither_count'] = counts.get('neither', 0)
                        iron_sponge_summary[f'{domain}_iron_percentage'] = percentages.get('iron', 0)
                        iron_sponge_summary[f'{domain}_sponge_percentage'] = percentages.get('sponge', 0)
                        iron_sponge_summary[f'{domain}_neither_percentage'] = percentages.get('neither', 0)
                
                # Individual question responses
                responses = iron_sponge_data.get('responses', {})
                for question_id, response_data in responses.items():
                    iron_sponge_questions[f'{question_id}_value'] = response_data.get('value', '')
                    iron_sponge_questions[f'{question_id}_domain'] = response_data.get('domain', '')
            
            # Build comprehensive data summary
            comprehensive_data = {
                # Basic info
                'participant_id': participant_id,
                'nickname': renew_data.get('nickname', ''),
                'started_at': renew_data.get('started_at', ''),
                
                # Completion status
                'completed_sections': ','.join(renew_data.get('completed_sections', [])),
                'sections_completed_count': len(renew_data.get('completed_sections', [])),
                
                # Response data
                'pos_recreation': renew_data.get('responses', {}).get('positive_events', {}).get('recreation', ''),
                'pos_achievement': renew_data.get('responses', {}).get('positive_events', {}).get('achievement', ''),
                'pos_relationship': renew_data.get('responses', {}).get('positive_events', {}).get('relationship', ''),
                'pos_other': renew_data.get('responses', {}).get('positive_events', {}).get('other', ''),
                'neg_adverse': renew_data.get('responses', {}).get('negative_events', {}).get('adverse', ''),
                'neg_neighbor': renew_data.get('responses', {}).get('negative_events', {}).get('neighbor', ''),
                'neg_discriminate': renew_data.get('responses', {}).get('negative_events', {}).get('discriminate', ''),
                'neg_stress': renew_data.get('responses', {}).get('negative_events', {}).get('stress', ''),
                
                # Journal entries
                'journal_things_to_change': journal_things_to_change,
                'journal_answer1': journal_answers.get('answer1', ''),
                'journal_answer2': journal_answers.get('answer2', ''),
                'journal_answer3': journal_answers.get('answer3', ''),
                'journal_answer4': journal_answers.get('answer4', ''),
                
                # Activity plan
                'posplan_who_category': posplan_responses.get('who', {}).get('category', ''),
                'posplan_who_response': posplan_responses.get('who', {}).get('response', ''),
                'posplan_where_category': posplan_responses.get('where', {}).get('category', ''),
                'posplan_where_response': posplan_responses.get('where', {}).get('response', ''),
                'posplan_what_activity': posplan_responses.get('what', {}).get('activity', ''),
                'posplan_when_timing': posplan_responses.get('when', {}).get('timing', ''),
                
                # Survey responses
                'feedback_statement1': feedback_survey.get('statement1', False),
                'feedback_statement2': feedback_survey.get('statement2', False),
                'feedback_statement3': feedback_survey.get('statement3', False),
                'feedback_statement4': feedback_survey.get('statement4', False),
                'journaling_preferences': ','.join(journaling_survey.get('responses', [])),
                'pmr_selected_options': ','.join(pmr_survey.get('selected_options', [])),
                'pmr_activity_id': pmr_survey.get('activity_id', ''),
                'posplan_sharing_preference': posplan_survey.get('sharing_preference', ''),
                'posplan_feeling': posplan_survey.get('feeling', ''),
                'posplan_sub_option': posplan_survey.get('sub_option', ''),
                
                # PMR sidebar survey data
                'pmr_selected_feelings': ','.join(pmr_sidebar_survey.get('selected_feelings', [])),
                
                # Journaling survey responses (more detailed)
                'journaling_survey_responses': json.dumps(journaling_survey),
                
                # Teaching survey data (entire object as JSON)
                'additional_teaching_survey': json.dumps(additional_teaching_survey),
                
                 # Individual privacy settings (entire object as JSON)
                'individual_privacy_settings': json.dumps(individual_privacy_settings),
                
                # Privacy settings
                'privacy_positive_events': renew_data.get('privacy_settings', {}).get('positive_events', False),
                'privacy_negative_events': renew_data.get('privacy_settings', {}).get('negative_events', False),
                'privacy_quiz_results': renew_data.get('privacy_settings', {}).get('quiz_results', False),
                'privacy_journal_responses': renew_data.get('privacy_settings', {}).get('journal_responses', False),
                'privacy_activity_plan': renew_data.get('privacy_settings', {}).get('activity_plan', False),
                
                # Iron & Sponge timestamp
                'iron_sponge_timestamp': iron_sponge_data.get('timestamp', ''),
            }
            
            # Add Iron & Sponge domain results
            comprehensive_data.update(iron_sponge_summary)
            
            # Add Iron & Sponge question responses
            comprehensive_data.update(iron_sponge_questions)
            
            # Calculate Iron & Sponge totals
            if iron_sponge_data.get('responses'):
                responses = iron_sponge_data['responses']
                total_iron = sum(1 for r in responses.values() if r.get('value') == 'iron')
                total_sponge = sum(1 for r in responses.values() if r.get('value') == 'sponge')
                total_neither = sum(1 for r in responses.values() if r.get('value') == 'neither')
                
                comprehensive_data.update({
                    'iron_sponge_total_iron_responses': total_iron,
                    'iron_sponge_total_sponge_responses': total_sponge,
                    'iron_sponge_total_neither_responses': total_neither
                })
            
            # Render the HTML summary for Qualtrics
            try:
                summary_html = render_to_string('accounts/qualtrics_summary.html', {
                    'nickname': renew_data.get('nickname', 'Anonymous'),
                    'participant_id': participant_id,
                    'sections_completed_count': len(renew_data.get('completed_sections', [])),
                    'positive_events': renew_data.get('responses', {}).get('positive_events', {}),
                    'negative_events': renew_data.get('responses', {}).get('negative_events', {}),
                    'journal_entries': journal_answers,
                    'activity_plan': posplan_responses,
                    'iron_sponge': iron_sponge_summary,
                })
            except Exception as e:
                print(f"Error rendering template: {e}")
                summary_html = f"<p>Error rendering summary for participant {participant_id}</p>"
            
            return JsonResponse({
                'success': True,
                'message': 'Data received for Qualtrics export',
                'participant_id': participant_id,
                'data_summary': {
                    'has_nickname': bool(renew_data.get('nickname')),
                    'sections_completed': len(renew_data.get('completed_sections', [])),
                    'total_notes': len(renew_data.get('notes', {})),
                    'has_quiz_results': bool(renew_data.get('quiz_results')),
                    'has_responses': bool(renew_data.get('responses')),
                    'has_journal_entries': bool(renew_data.get('journal_entries')),
                    'has_posplan_responses': bool(renew_data.get('posplan_responses')),
                    'has_iron_sponge_results': bool(iron_sponge_data),
                    'iron_sponge_questions_answered': len(iron_sponge_data.get('responses', {}))
                },
                'qualtrics_data': comprehensive_data,  # Structured data
                'summary_html': summary_html  # HTML summary for display
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Invalid method'}, status=405)


@csrf_exempt
def save_iron_sponge(request):
    """Save Iron & Sponge quiz results to the database (keyed by session_key).
    Expects JSON: { "ironSpongeResults": { school, friend, family, responses, nickname } }
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body or b'{}')
            iron = data.get('ironSpongeResults') or data.get('iron_sponge_results') or {}
            nickname = data.get('nickname') or iron.get('nickname', '')

            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            def _counts(domain):
                d = iron.get(domain, {}) or {}
                c = d.get('counts', {}) or {}
                return d.get('type', 'balanced'), c.get('iron', 0), c.get('sponge', 0), c.get('neither', 0)

            sch_type, sch_iron, sch_sponge, sch_neither = _counts('school')
            fri_type, fri_iron, fri_sponge, fri_neither = _counts('friend')
            fam_type, fam_iron, fam_sponge, fam_neither = _counts('family')

            _purge_other_sessions(session_key)

            from accounts.models import IronSpongeResponse
            IronSpongeResponse.objects.update_or_create(
                session_key=session_key,
                defaults={
                    'nickname': nickname,
                    'school_type': sch_type, 'school_iron': sch_iron, 'school_sponge': sch_sponge, 'school_neither': sch_neither,
                    'friend_type': fri_type, 'friend_iron': fri_iron, 'friend_sponge': fri_sponge, 'friend_neither': fri_neither,
                    'family_type': fam_type, 'family_iron': fam_iron, 'family_sponge': fam_sponge, 'family_neither': fam_neither,
                    'responses_json': iron.get('responses', {}),
                    'user': request.user if request.user.is_authenticated else None,
                }
            )

            return JsonResponse({'success': True, 'session_key': session_key})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)
    
@csrf_exempt
def save_posplan_activity(request):
    """Save Positive Planned Activity responses (Who/Where/What/When) to the database."""
    if request.method == 'POST':
        try:
            data = json.loads(request.body or b'{}')
            nickname = data.get('nickname', '')
            who = data.get('who', '')
            where = data.get('where', '')
            what = data.get('what', '')
            when = data.get('when', '')

            if not request.session.session_key:
                request.session.create()
            session_key = request.session.session_key

            _purge_other_sessions(session_key)

            from accounts.models import PosplanActivityResponse
            PosplanActivityResponse.objects.update_or_create(
                session_key=session_key,
                defaults={
                    'nickname': nickname,
                    'who': who,
                    'where': where,
                    'what': what,
                    'when': when,
                    'user': request.user if request.user.is_authenticated else None,
                }
            )
            return JsonResponse({'success': True, 'session_key': session_key})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid method'}, status=405)


def qualtrics_survey_view(request):
    """Display embedded Qualtrics survey"""
    return render(request, 'accounts/qualtrics_survey.html')
    
def completion_video(request):
    return render(request, 'accounts/completion_video.html')


def _build_beautiful_summary_context(session_key):
    """Build context for the redesigned summary template from session-linked DB data."""
    from accounts.models import (
        PositiveEventResponse,
        NegativeEventResponse,
        PMRSurveyResponse,
        IronSpongeResponse,
        PosplanActivityResponse,
        FeedbackSurveyResponse,
        AttentionCheckResponse,
        JournalEntry,
    )

    pos_responses = list(PositiveEventResponse.objects.filter(session_key=session_key))
    neg_responses = list(NegativeEventResponse.objects.filter(session_key=session_key))
    pmr_response = PMRSurveyResponse.objects.filter(session_key=session_key).first()
    iron_sponge_db = IronSpongeResponse.objects.filter(session_key=session_key).first()
    posplan_db = PosplanActivityResponse.objects.filter(session_key=session_key).first()
    feedback_db = FeedbackSurveyResponse.objects.filter(session_key=session_key).first()
    attention_checks = list(AttentionCheckResponse.objects.filter(session_key=session_key).order_by('question_id'))

    positive_events = {r.category: r.response_text for r in pos_responses}
    negative_events = {r.category: r.response_text for r in neg_responses}

    nickname = 'Participant'
    for candidate in [
        (pos_responses[0].nickname if pos_responses else None),
        (neg_responses[0].nickname if neg_responses else None),
        (pmr_response.nickname if pmr_response else None),
        (iron_sponge_db.nickname if iron_sponge_db else None),
        (posplan_db.nickname if posplan_db else None),
        (feedback_db.nickname if feedback_db else None),
    ]:
        if candidate:
            nickname = candidate
            break

    activity_plan = {
        'who': {'response': '', 'category': ''},
        'where': {'response': '', 'category': ''},
        'what': {'activity': ''},
        'when': {'timing': ''},
    }
    if posplan_db:
        activity_plan = {
            'who': {'response': posplan_db.who or '', 'category': ''},
            'where': {'response': posplan_db.where or '', 'category': ''},
            'what': {'activity': posplan_db.what or ''},
            'when': {'timing': posplan_db.when or ''},
        }

    iron_sponge = {
        'school_type': '',
        'school_iron_count': 0,
        'school_sponge_count': 0,
        'school_neither_count': 0,
        'family_type': '',
        'family_iron_count': 0,
        'family_sponge_count': 0,
        'family_neither_count': 0,
        'friend_type': '',
        'friend_iron_count': 0,
        'friend_sponge_count': 0,
        'friend_neither_count': 0,
    }
    if iron_sponge_db:
        iron_sponge = {
            'school_type': iron_sponge_db.school_type or '',
            'school_iron_count': iron_sponge_db.school_iron,
            'school_sponge_count': iron_sponge_db.school_sponge,
            'school_neither_count': iron_sponge_db.school_neither,
            'family_type': iron_sponge_db.family_type or '',
            'family_iron_count': iron_sponge_db.family_iron,
            'family_sponge_count': iron_sponge_db.family_sponge,
            'family_neither_count': iron_sponge_db.family_neither,
            'friend_type': iron_sponge_db.friend_type or '',
            'friend_iron_count': iron_sponge_db.friend_iron,
            'friend_sponge_count': iron_sponge_db.friend_sponge,
            'friend_neither_count': iron_sponge_db.friend_neither,
        }

    # Build implementation feedback context
    pmr_feedback = {}
    if pmr_response and pmr_response.selected_feelings:
        pmr_feedback = {
            'question': f"How did {nickname} body feel after PMR?",
            'response': ', '.join(pmr_response.selected_feelings),
            'selected': list(pmr_response.selected_feelings),
        }
    else:
        pmr_feedback = {'question': f"How did {nickname} body feel after PMR?", 'response': '', 'selected': []}

    thoughts_question_ids = ['emo_p2_q1', 'emo_p2_q2', 'emo_p2_q3', 'emo_p2_q4']
    thoughts_trigger_labels = [
        'I liked learning about how the world around us affects people, and I want to talk to someone more about this.',
        'I liked learning about how the world around us affects people, and I want to see or read more about it.',
        'To be honest, I found some of this stuff a little boring, but I want to talk more about thoughts, feelings, and emotions.',
        "To be honest, I didn't find what we talked about so far very interesting or helpful.",
    ]

    thoughts_correct_count = sum(
        1 for check in attention_checks
        if check.question_id in thoughts_question_ids and check.answered_correctly
    )
    thoughts_part4_done = 0
    if feedback_db and feedback_db.responses:
        saved_feedback = feedback_db.responses or []
        if any(label in saved_feedback for label in thoughts_trigger_labels):
            thoughts_part4_done = 1

    thoughts_progress = min((thoughts_correct_count + thoughts_part4_done) / 5.0, 1.0)

    implementation_feedback = {}
    if feedback_db and feedback_db.responses:
        # responses is a list of selected feedback statements
        for response in feedback_db.responses:
            response_lower = response.lower()
            if 'psychoeducation' in response_lower or 'feel about psychoeducation' in response_lower:
                implementation_feedback['psychoeducation'] = response
            elif 'body' in response_lower and 'pmr' in response_lower:
                implementation_feedback['pmr_feeling'] = response
            elif 'again' in response_lower or 'try it again' in response_lower:
                implementation_feedback['pmr_again'] = response
            elif 'journaling' in response_lower and 'feel about journaling' in response_lower:
                implementation_feedback['journaling'] = response
            elif 'optimistic' in response_lower or 'activity plan' in response_lower:
                if 'optimistic' in response_lower:
                    implementation_feedback['activity_plan_optimistic'] = response
            elif 'support' in response_lower and 'journaling' in response_lower:
                implementation_feedback['journaling_support'] = response

    # If responses weren't properly categorized, check if we have a single psychoeducation response
    if not implementation_feedback and feedback_db and feedback_db.responses:
        if len(feedback_db.responses) > 0:
            implementation_feedback['psychoeducation'] = feedback_db.responses[0]

    # Determine PMR 'try again' flag from implementation feedback if available
    pmr_wants_try = False
    pmr_try_text = ''
    pmr_try_list = []

    # Known PMR option labels (must match texts in pmr_p1.html)
    pmr_option_labels = [
        'I liked it',
        'I want to try it on my own',
        'I would like to try doing it more with a counselor',
        "I didn't like it",
        'I would prefer not to try something like that again',
        'I would be open to trying it with a counselor again to see if there was something I was doing wrong',
        'I would be interested in talking to a counselor about other things I can practice to feel more relaxed',
        "I don't know how I feel about it",
    ]

    if feedback_db and feedback_db.responses:
        # Build a selection list for PMR options based on stored responses
        responses = feedback_db.responses or []
        for opt in pmr_option_labels:
            pmr_try_list.append({'label': opt, 'selected': opt in responses})
        # If any PMR-related option selected, set pmr_wants_try
        if any(item['selected'] for item in pmr_try_list):
            pmr_wants_try = True
            pmr_try_text = ', '.join([item['label'] for item in pmr_try_list if item['selected']])
    else:
        # fallback to previous implementation_feedback key if present
        if implementation_feedback.get('pmr_again'):
            pmr_wants_try = True
            pmr_try_text = implementation_feedback.get('pmr_again')
            # also include as simple list
            pmr_try_list = [{'label': pmr_try_text, 'selected': True}]

    # Build psychoeducation feedback list similar to pmr_try_list
    psychoeducation_feedback_list = []
    
    # Known psychoeducation option labels
    psychoeducation_option_labels = [
        'I liked learning about how the world around us affects people, and I want to talk to someone more about this.',
        'I liked learning about how the world around us affects people, and I want to see or read more about it.',
        'To be honest, I found some of this stuff a little boring, but I want to talk more about thoughts, feelings, and emotions.',
        "To be honest, I didn't find what we talked about so far very interesting or helpful.",
    ]
    
    if feedback_db and feedback_db.responses:
        # Build a selection list for psychoeducation options based on stored responses
        responses = feedback_db.responses or []
        for opt in psychoeducation_option_labels:
            psychoeducation_feedback_list.append({'label': opt, 'selected': opt in responses})

    # Build attention check context in the same order as the worksheet screenshot
    attention_check_data = []
    question_text_map = {
        'neg_p6_q1': 'Why might Mike react more to Coach Ramos\'s comments than Jackie does?',
        'neg_p7_q1': 'How might Mike be feeling after Coach Ramos\'s comment?',
        'emo_p2_q1': 'Dave is sad because he has had angry thoughts about the world from a very young age.',
        'emo_p2_q2': 'Avoiding friends, like Dave is, is a common for youth who had a bad thing happen to them like losing a loved one.',
        'emo_p2_q3': 'It\'s strange that Dave is fighting with his mom, as kids are usually just avoidant after a major life event.',
        'emo_p2_q4': 'Dave may keep thinking about his grandfather, and feeling sad in the room, because he is being "triggered".',
    }
    # Explicitly group Mike items first, then Dave items, to match the worksheet layout.
    attention_order = ['neg_p6_q1', 'neg_p7_q1', 'emo_p2_q1', 'emo_p2_q2', 'emo_p2_q3', 'emo_p2_q4']
    attention_lookup = {check.question_id: check for check in attention_checks}

    for question_id in attention_order:
        check = attention_lookup.get(question_id)
        question_data = {
            'question_id': question_id,
            'text': question_text_map.get(question_id, ''),
            'attempts': 0,
            'answered': False,
        }
        if check:
            question_data['attempts'] = check.incorrect_attempts + 1  # +1 to include the correct attempt
            question_data['answered'] = True
        attention_check_data.append(question_data)

    # Fetch journal entries from DB
    journal_db = JournalEntry.objects.filter(session_key=session_key).first()
    journal_entries = {}
    if journal_db:
        journal_entries = {
            'things_to_change': journal_db.things_to_change,
            'answer1': journal_db.experiences_shaped_me,  # map to answer1 for compatibility
            'raw_content': journal_db.experiences_shaped_me,
        }

    # Build journaling feedback list from any saved FeedbackSurveyResponse.responses
    journaling_feedback_list = []
    journaling_option_labels = [
        'I like writing about these things, and I want to share it with a family member or friend.',
        'I like writing about these things, and I want to share it with a counselor.',
        "I like writing about these things, but I don't want to share with anyone.",
        "I don't like writing about these things. I would rather talk to someone.",
        "I don't like writing or talking about these things. It makes me feel worse.",
    ]
    if feedback_db and feedback_db.responses:
        saved = feedback_db.responses or []
        for label in journaling_option_labels:
            journaling_feedback_list.append({'label': label, 'selected': label in saved})

    journaling_part1_progress = 0.0
    journaling_part2_progress = 0.0
    journaling_part3_progress = 0.0

    # Journaling part 1: Things I Want to Change
    if journal_db and (journal_db.things_to_change or '').strip():
        journaling_part1_progress = 1.0

    # Journaling part 2: Experiences That Shaped Me
    if journal_db and (journal_db.experiences_shaped_me or '').strip():
        journaling_part2_progress = 1.0

    # Journaling part 3: Journaling Feedback modal submit
    if feedback_db and feedback_db.responses:
        saved = feedback_db.responses or []
        if any(label in saved for label in journaling_option_labels):
            journaling_part3_progress = 1.0

    # Journaling is split into 3 submitted parts.
    journaling_progress = (journaling_part1_progress + journaling_part2_progress + journaling_part3_progress) / 3.0

    # Build Positive Activity Plan feeling list from FeedbackSurveyResponse.responses
    activity_plan_list = []
    activity_option_labels = [
        'I think this positive planned activity will happen',
        "I'm excited",
        'But I\'d love to talk to a counselor about it to make sure it happens!',
        "I don't think this will happen",
        "I'm not sure this will happen",
        "But I'd like to talk to a counselor to give me help to make it happen",
        'I would like to share it with someone else at school',
        "I don't want to share this with anyone",
        "But I'd be open to talking to a counselor about it",
        "I don't want to talk to a counselor about it",
    ]
    if feedback_db and feedback_db.responses:
        saved = feedback_db.responses or []
        for label in activity_option_labels:
            activity_plan_list.append({'label': label, 'selected': label in saved})

    pmr_part1_progress = 1.0 if (pmr_response and pmr_response.selected_feelings) else 0.0
    pmr_part2_progress = 0.0
    if feedback_db and feedback_db.responses:
        pmr_saved = feedback_db.responses or []
        if any(label in pmr_saved for label in pmr_option_labels):
            pmr_part2_progress = 1.0

    # PMR is split into two submitted parts; each part contributes half of PMR's skill weight.
    pmr_skill_progress = (pmr_part1_progress + pmr_part2_progress) / 2.0

    # Positive Planned Activities: split into 5 parts - who, where, what, when, and feeling
    posplan_part_who = 1.0 if (posplan_db and (posplan_db.who or '').strip()) else 0.0
    posplan_part_where = 1.0 if (posplan_db and (posplan_db.where or '').strip()) else 0.0
    posplan_part_what = 1.0 if (posplan_db and (posplan_db.what or '').strip()) else 0.0
    posplan_part_when = 1.0 if (posplan_db and (posplan_db.when or '').strip()) else 0.0

    posplan_part5_progress = 0.0
    if feedback_db and feedback_db.responses:
        saved = feedback_db.responses or []
        if any(label in saved for label in activity_option_labels):
            posplan_part5_progress = 1.0

    posplan_skill_progress = (posplan_part_who + posplan_part_where + posplan_part_what + posplan_part_when + posplan_part5_progress) / 5.0
    skill_progress = (pmr_skill_progress + journaling_progress + posplan_skill_progress) / 3.0

    # Compute per-section fractional progress (0..1) so overall percent is accurate.
    # Positive/Negative: fraction = answered categories / 4
    # Thoughts: completed trigger set => full
    # Skills: PMR, Journaling, and Positive Planned Activities each contribute one-third
    try:
        positive_progress = min(len(positive_events) / 4.0, 1.0)
    except Exception:
        positive_progress = 0.0

    # Negative events are composed of the four interactive cards, the two attention-check
    # quiz questions, and the Iron & Sponge quiz. Cards contribute 50% of the section,
    # the two quiz questions contribute 25%, and Iron & Sponge contributes 25%.
    try:
        cards_frac = min(len(negative_events) / 4.0, 1.0)
    except Exception:
        cards_frac = 0.0

    attention_check_ids = {'neg_p6_q1', 'neg_p7_q1'}
    attention_check_completed = 0
    for check in attention_checks:
        if check.question_id in attention_check_ids and check.answered_correctly:
            attention_check_completed += 1
    attention_frac = min(attention_check_completed / 2.0, 1.0)

    quiz_frac = 0.0
    if iron_sponge_db and iron_sponge_db.responses_json:
        try:
            quiz_frac = min(len(iron_sponge_db.responses_json.keys()) / 9.0, 1.0)
        except Exception:
            quiz_frac = 1.0
    negative_progress = (cards_frac * 0.5) + (attention_frac * 0.25) + (quiz_frac * 0.25)

    # Count only fully completed sections for the 'Sections complete: X/4' label
    sections_completed_count = sum([
        1 if positive_progress >= 1.0 else 0,
        1 if negative_progress >= 1.0 else 0,
        1 if thoughts_progress >= 1.0 else 0,
        1 if skill_progress >= 1.0 else 0,
    ])

    # Overall percent is the average of the four section progresses (0..100)
    overall_percent = int(round(((positive_progress + negative_progress + thoughts_progress + skill_progress) / 4.0) * 100))

    return {
        'participant_id': session_key,
        'nickname': nickname,
        'sections_completed_count': sections_completed_count,
        'overall_percent': overall_percent,
        'positive_events': positive_events,
        'negative_events': negative_events,
        'pmr_feedback': pmr_feedback,
        'pmr_wants_try': pmr_wants_try,
        'pmr_try_text': pmr_try_text,
        'pmr_try_list': pmr_try_list,
        'psychoeducation_feedback_list': psychoeducation_feedback_list,
        'journal_entries': journal_entries,
        'activity_plan': activity_plan,
        'iron_sponge': iron_sponge,
        'implementation_feedback': implementation_feedback,
        'attention_check': attention_check_data,
        'journaling_feedback_list': journaling_feedback_list,
        'activity_plan_list': activity_plan_list,
        'current_session_key': session_key,
    }


def beautiful_summary_view(request, session_key=None):
    """Render redesigned summary sheet in a standalone, print-friendly page."""
    if not session_key:
        session_key = request.session.session_key
    context = _build_beautiful_summary_context(session_key or '')
    return render(request, 'accounts/summary_download_page.html', context)


def download_beautiful_summary(request, session_key=None):
    """Download redesigned summary as a standalone HTML file."""
    if not session_key:
        session_key = request.session.session_key
    context = _build_beautiful_summary_context(session_key or '')
    html = render_to_string('accounts/summary_download_page.html', context, request=request)
    safe_name = ''.join(c if c.isalnum() or c in ('-', '_') else '_' for c in context['nickname']).strip('_') or 'participant'

    response = HttpResponse(html, content_type='text/html; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="RENEW-Beautiful-Summary-{safe_name}.html"'
    return response


def download_summary(request, session_key=None):
    """Generate and download a filled RENEW Summary Sheet as a .docx file.
    
    Usage:
        /download-summary/                          → uses the current user's session
        /download-summary/<session_key>/             → uses a specific session
    """
    from docx import Document
    from accounts.models import PositiveEventResponse, NegativeEventResponse, PMRSurveyResponse, AttentionCheckResponse, IronSpongeResponse, PosplanActivityResponse

    # Determine which session to pull data for
    if not session_key:
        session_key = request.session.session_key

    if not session_key:
        return HttpResponse('No session found. Please complete an activity first.', status=404)

    pos_responses = PositiveEventResponse.objects.filter(session_key=session_key)
    neg_responses = NegativeEventResponse.objects.filter(session_key=session_key)
    pmr_response = PMRSurveyResponse.objects.filter(session_key=session_key).first()

    # Also allow generation if Iron & Sponge results were saved in the session
    iron_sponge_db = IronSpongeResponse.objects.filter(session_key=session_key).first()
    posplan_db = PosplanActivityResponse.objects.filter(session_key=session_key).first()

    if not pos_responses.exists() and not neg_responses.exists() and not pmr_response and not iron_sponge_db and not posplan_db:
        return HttpResponse('No responses found for this session.', status=404)

    # Build a dict of category → response_text for positive events
    nickname = 'Unknown'
    response_map = {}
    for r in pos_responses:
        response_map[r.category] = r.response_text
        if r.nickname:
            nickname = r.nickname

    # Build a dict of category → response_text for negative events
    neg_response_map = {}
    for r in neg_responses:
        neg_response_map[r.category] = r.response_text
        if r.nickname and nickname == 'Unknown':
            nickname = r.nickname

    # Get nickname from PMR if still unknown
    if pmr_response and pmr_response.nickname and nickname == 'Unknown':
        nickname = pmr_response.nickname

    # Load the template .docx
    template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'RENEW Summary Sheet.docx')
    doc = Document(template_path)

    # Category → paragraph index mapping
    category_to_paragraph = {
        'recreation': 2,
        'achievement': 3,
        'relationship': 4,
        'other': 5,
    }

    # Replace (name) with the nickname throughout
    for para in doc.paragraphs:
        if '(name)' in para.text:
            _replace_in_paragraph(para, '(name)', nickname)

    # Fill in the answers: remove placeholder labels from Q1 & Q2,
    # then append all answers on a new line in italics (same style for all 4).
    placeholder_labels = {
        'recreation': ('(Recreational positive event)', 2),
        'achievement': ('(achievement positive event)', 3),
    }
    # Strip the parenthetical placeholders from Q1 & Q2
    for category, (label, para_idx) in placeholder_labels.items():
        _replace_in_paragraph(doc.paragraphs[para_idx], f' {label}', '')

    append_categories = {
        'recreation': 2,
        'achievement': 3,
        'relationship': 4,
        'other': 5,
    }

    # Process from bottom to top so inserted spacing does not shift later target paragraphs.
    for category, para_idx in sorted(append_categories.items(), key=lambda item: item[1], reverse=True):
        answer = response_map.get(category, '')
        if answer:
            from docx.shared import RGBColor

            run = doc.paragraphs[para_idx].add_run(f'\n    {answer}')
            run.italic = True
            run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Negative events: append answers to P10-P13 ---
    neg_append_categories = {
        'adverse': 10,
        'neighbor': 11,
        'discriminate': 12,
        'stress': 13,
    }

    for category, para_idx in neg_append_categories.items():
        answer = neg_response_map.get(category, '')
        if answer:
            from docx.shared import RGBColor

            run = doc.paragraphs[para_idx].add_run(f'\n    {answer}')
            run.italic = True
            run.font.color.rgb = RGBColor(192, 0, 0)

    # --- PMR survey: append feelings to P20 ---
    if pmr_response and pmr_response.selected_feelings:
        feelings_text = ', '.join(pmr_response.selected_feelings)
        from docx.shared import RGBColor
        run = doc.paragraphs[20].add_run(f'\n    {feelings_text}')
        run.italic = True
        run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Positive Planned Activity: read from DB and append to P36 ---
    if posplan_db:
        lines = []
        if posplan_db.who:
            lines.append(f'Who: {posplan_db.who}')
        if posplan_db.where:
            lines.append(f'Where: {posplan_db.where}')
        if posplan_db.what:
            lines.append(f'What: {posplan_db.what}')
        if posplan_db.when:
            lines.append(f'When: {posplan_db.when}')
        if lines:
            from docx.shared import RGBColor
            run = doc.paragraphs[36].add_run('\n    ' + '\n    '.join(lines))
            run.italic = True
            run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Iron & Sponge: read from DB and append to P35 ---
    if iron_sponge_db:
        try:
            domain_lines = [
                f"School: {(iron_sponge_db.school_type or 'balanced').capitalize()} ({iron_sponge_db.school_iron} Iron, {iron_sponge_db.school_sponge} Sponge, {iron_sponge_db.school_neither} Neither)",
                f"Friend: {(iron_sponge_db.friend_type or 'balanced').capitalize()} ({iron_sponge_db.friend_iron} Iron, {iron_sponge_db.friend_sponge} Sponge, {iron_sponge_db.friend_neither} Neither)",
                f"Family: {(iron_sponge_db.family_type or 'balanced').capitalize()} ({iron_sponge_db.family_iron} Iron, {iron_sponge_db.family_sponge} Sponge, {iron_sponge_db.family_neither} Neither)",
            ]
            from docx.shared import RGBColor
            run = doc.paragraphs[35].add_run('\n    ' + '\n    '.join(domain_lines))
            run.italic = True
            run.font.color.rgb = RGBColor(192, 0, 0)
        except Exception:
            pass

    # --- Attention checks: append incorrect attempt counts to P28-P33 ---
    attention_responses = AttentionCheckResponse.objects.filter(session_key=session_key)
    attention_map = {r.question_id: r for r in attention_responses}

    # question_id → (paragraph index, total_attempts_formula)
    attention_questions = {
        'neg_p6_q1': 28,
        'neg_p7_q1': 29,
        'emo_p2_q1': 30,
        'emo_p2_q2': 31,
        'emo_p2_q3': 32,
        'emo_p2_q4': 33,
    }

    for qid, para_idx in attention_questions.items():
        resp = attention_map.get(qid)
        if resp:
            total_attempts = resp.incorrect_attempts + (1 if resp.answered_correctly else 0)
            from docx.shared import RGBColor
            run = doc.paragraphs[para_idx].add_run(f'  →  {total_attempts} attempt(s)')
            run.italic = True
            run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Feedback survey (emo_p4) responses: append summary to P22 ---
    try:
        from accounts.models import FeedbackSurveyResponse
        # Use the feedback response for this session so answer changes are reflected per user/session.
        feedback = FeedbackSurveyResponse.objects.filter(session_key=session_key).order_by('-created_at').first()
        if feedback:
            # Primary source: the JSON `responses` column.
            lines = list(feedback.responses or [])

            # Backward-compatible fallback if `responses` is empty on older rows.
            if not lines:
                if feedback.statement1_checked:
                    lines.append('I liked learning about how the world around us affects people, and I want to talk to someone more about this.')
                if feedback.statement2_checked:
                    lines.append('I liked learning about how the world around us affects people, and I want to see or read more about it.')
                if feedback.statement3_checked:
                    lines.append('To be honest, I found some of this stuff a little boring, but I want to talk more about thoughts, feelings, and emotions.')
                if feedback.statement4_checked:
                    lines.append('To be honest, I didn\'t find what we talked about so far very interesting or helpful.')

            if lines:
                # Find the psychoeducation question paragraph dynamically so template edits don't break this.
                feedback_para_idx = None
                for idx, para in enumerate(doc.paragraphs):
                    para_text = (para.text or '').lower()
                    if 'how did' in para_text and 'psychoeducation' in para_text:
                        feedback_para_idx = idx
                        break

                if feedback_para_idx is None:
                    feedback_para_idx = 18  # Safe fallback for current template

                from docx.shared import RGBColor

                # Render each selected response on its own clean line under the question.
                formatted_lines = '\n'.join(f'    {line}' for line in lines)
                run = doc.paragraphs[feedback_para_idx].add_run('\n' + formatted_lines)
                run.italic = True
                run.font.color.rgb = RGBColor(192, 0, 0)
    except Exception:
        pass

    # Write to an in-memory buffer and return as a download
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f'RENEW Summary Sheet - {nickname}.docx'
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

