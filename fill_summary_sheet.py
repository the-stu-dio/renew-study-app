"""
Fill the RENEW Summary Sheet.docx with data from the database.

Usage:
    python fill_summary_sheet.py [session_key]

If no session_key is provided, it uses the most recent session's data.
The filled document is saved as "RENEW Summary Sheet - <nickname>.docx".
"""

import sqlite3
import sys
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DB_PATH = os.path.join(os.path.dirname(__file__), 'db.sqlite3')
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'RENEW Summary Sheet.docx')


def get_positive_responses(session_key=None):
    """Fetch positive event responses from the database."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    if session_key:
        cursor.execute(
            "SELECT nickname, category, response_text FROM accounts_positiveeventresponse WHERE session_key = ?",
            (session_key,)
        )
    else:
        # Get the most recent session_key
        cursor.execute(
            "SELECT session_key FROM accounts_positiveeventresponse ORDER BY created_at DESC LIMIT 1"
        )
        row = cursor.fetchone()
        if not row:
            print("No positive event responses found in the database.")
            sys.exit(1)
        session_key = row['session_key']
        cursor.execute(
            "SELECT nickname, category, response_text FROM accounts_positiveeventresponse WHERE session_key = ?",
            (session_key,)
        )

    rows = cursor.fetchall()
    conn.close()

    if not rows:
        print(f"No responses found for session_key: {session_key}")
        sys.exit(1)

    nickname = rows[0]['nickname'] or 'Unknown'
    responses = {row['category']: row['response_text'] for row in rows}
    return nickname, responses, session_key


def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph while preserving formatting."""
    if old_text not in paragraph.text:
        return False

    # Try run-level replacement first (preserves formatting)
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # If the target text spans multiple runs, rebuild the paragraph
    full_text = paragraph.text
    if old_text in full_text:
        new_full_text = full_text.replace(old_text, new_text)
        # Clear all runs except the first, put all text in the first run
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = new_full_text
            else:
                run.text = ''
        return True

    return False


def append_response_below_question(paragraph, response_text):
    """Add the response on its own line with extra spacing and red text."""
    run = paragraph.add_run(f'\n    {response_text}')
    run.italic = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    return run


def fill_summary_sheet(nickname, responses):
    """Fill the Word document template with the provided data."""
    doc = Document(TEMPLATE_PATH)

    # Category → paragraph index mapping (from the doc inspection):
    # P2: "What does (name) like to do for fun? (Recreational positive event)?"  → recreation
    # P3: "What kinds of events give (name) confidence? (achievement positive event)?" → achievement
    # P4: "Who are the people that are most positive and important people to (name)?" → relationship
    # P5: "Here are some additional positive things (name) mentioned..." → other

    category_to_paragraph = {
        'recreation': 2,
        'achievement': 3,
        'relationship': 4,
        'other': 5,
    }

    # First, replace (name) with the nickname in ALL paragraphs
    for para in doc.paragraphs:
        if '(name)' in para.text:
            replace_in_paragraph(para, '(name)', nickname)

    # Now append the response text to each relevant paragraph.
    # Process from bottom to top so inserting blank lines does not shift later targets.
    for category, para_idx in sorted(category_to_paragraph.items(), key=lambda item: item[1], reverse=True):
        response = responses.get(category, '')
        if response:
            para = doc.paragraphs[para_idx]
            # Put the response on its own line with extra space and red text
            append_response_below_question(para, response)

    # Save the filled document
    output_filename = f'RENEW Summary Sheet - {nickname}.docx'
    output_path = os.path.join(os.path.dirname(__file__), output_filename)
    doc.save(output_path)
    return output_path


def main():
    session_key = sys.argv[1] if len(sys.argv) > 1 else None

    print("Fetching data from database...")
    nickname, responses, session_key = get_positive_responses(session_key)

    print(f"  Session:  {session_key}")
    print(f"  Nickname: {nickname}")
    for cat, text in responses.items():
        print(f"  {cat}: {text}")

    print("\nFilling summary sheet...")
    output_path = fill_summary_sheet(nickname, responses)
    print(f"Done! Saved to: {output_path}")


if __name__ == '__main__':
    main()
